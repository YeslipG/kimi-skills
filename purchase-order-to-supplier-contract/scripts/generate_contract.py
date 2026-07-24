#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据采购订单/销售合同（PDF 或图片）+ Word 模板，自动生成供应商侧《工矿机械产品买卖合同》Word 文档。

规则：
- 新合同签订日期 = 文件1签订/订单日期 + 5 天
- 交（提）货时间 = 文件1签订/订单日期 + 60 天
- 单价、总价按文件1含税金额 × 折扣（默认 0.8）填写
- 合同编号 = HPGH + 签订日期(yy/mm/dd)简写 + "-1"
- 结算条款中 BGJ 开头的销售合同号，取自文件1文件名；无则替换为"补充销售合同号"
- 文件名 = 当日日期(YYMMDD)-hpgh-买方全称-折后总价.doc
"""

import argparse
import json
import os
import re
import shutil
import sys
from copy import deepcopy
from datetime import datetime, timedelta
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt, Cm
from docx.table import _Row

# ---------------------------------------------------------------------------
# 数字金额转中文大写
# ---------------------------------------------------------------------------
CN_NUM = "零壹贰叁肆伍陆柒捌玖"
CN_UNIT = ["", "拾", "佰", "仟"]
CN_BIG_UNIT = ["", "万", "亿", "万亿"]


def _int_to_chinese(n: int) -> str:
    if n == 0:
        return "零"
    parts = []
    group_idx = 0
    while n > 0:
        group = n % 10000
        n //= 10000
        if group == 0:
            if parts:
                parts.append(CN_BIG_UNIT[group_idx])
            group_idx += 1
            continue
        group_str = ""
        zero = False
        for i in range(4):
            digit = group % 10
            group //= 10
            if digit == 0:
                if not zero and group_str:
                    zero = True
            else:
                if zero:
                    group_str = "零" + group_str
                    zero = False
                group_str = CN_NUM[digit] + CN_UNIT[i] + group_str
        parts.append(group_str + CN_BIG_UNIT[group_idx])
        group_idx += 1
    return "".join(reversed(parts))


def money_to_chinese(amount: float) -> str:
    """把金额（元）转为中文大写，含角分。"""
    yuan = int(amount)
    jiao_fen = round((amount - yuan) * 100)
    jiao = jiao_fen // 10
    fen = jiao_fen % 10
    result = _int_to_chinese(yuan) + "元"
    if jiao == 0 and fen == 0:
        result += "整"
    else:
        if jiao > 0:
            result += CN_NUM[jiao] + "角"
        if fen > 0:
            result += CN_NUM[fen] + "分"
    return result


# ---------------------------------------------------------------------------
# 文本替换辅助
# ---------------------------------------------------------------------------
def replace_in_paragraph(p, old: str, new: str) -> bool:
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    idx = full.find(old)
    end = idx + len(old)
    pos = 0
    done = False
    for r in p.runs:
        rlen = len(r.text)
        rstart, rend = pos, pos + rlen
        if rend > idx and rstart < end:
            head = r.text[: max(0, idx - rstart)]
            tail = r.text[max(0, end - rstart) :]
            r.text = head + (new if not done else "") + tail
            done = True
        pos = rend
    return True


def replace_in_cell(cell, old: str, new: str) -> bool:
    changed = False
    for p in cell.paragraphs:
        if replace_in_paragraph(p, old, new):
            changed = True
    return changed


def set_cell_text(cell, text: str):
    """清空单元格并写入新文本（保留第一个段落）。"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if cell.paragraphs:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = text
        else:
            cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)


# ---------------------------------------------------------------------------
# 买方简称
# ---------------------------------------------------------------------------
def short_buyer(name: str) -> str:
    for suffix in ["股份有限公司", "有限责任公司", "有限公司"]:
        if name.endswith(suffix):
            return name[: -len(suffix)]
    return name


# ---------------------------------------------------------------------------
# 文件名中提取销售合同号
# ---------------------------------------------------------------------------
def extract_sales_contract_no(filename: str) -> str:
    """从文件1文件名末尾提取 BGJ 开头、数字结尾的合同号；无则返回"补充销售合同号"。"""
    base = Path(filename).stem
    m = re.search(r"BGJ[^\\/\\s]*?\d+(?=\.[^.]*$|$)", base)
    if m:
        return m.group(0)
    return "补充销售合同号"


# ---------------------------------------------------------------------------
# PDF 信息提取
# ---------------------------------------------------------------------------
def parse_date(text: str):
    for pat in [r"签订时间[：:]\s*(\d{4}-\d{2}-\d{2})", r"签订日期[：:]\s*(\d{4}-\d{2}-\d{2})",
                r"订单日期[：:]\s*(\d{4}-\d{2}-\d{2})", r"(\d{4}年\d{2}月\d{2}日)"]:
        m = re.search(pat, text)
        if m:
            s = m.group(1)
            if "年" in s:
                return datetime.strptime(s, "%Y年%m月%d日").date()
            return datetime.strptime(s, "%Y-%m-%d").date()
    return None


def extract_pdf_info(pdf_path: str):
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[0]
        text = page.extract_text() or ""
        tables = page.extract_tables()

    info = {
        "buyer": "",
        "seller": "",
        "order_date": parse_date(text),
        "items": [],
        "note": "",
    }

    # 买方/卖方：支持多种表述
    for pattern in [
        r"买受方[：:]\s*(.+?)(?:\s|签订|$)",
        r"买受人(?:\(甲方\)|（甲方）)?[：:]\s*(.+?)(?:\s|签订|出卖人|$)",
        r"甲方(?:\(买受方\)|（买受方）)?[：:]\s*(.+?)(?:\s+乙方|$)",
    ]:
        m = re.search(pattern, text)
        if m:
            info["buyer"] = m.group(1).strip()
            break

    for pattern in [
        r"出卖方[：:]\s*(.+?)(?:\s|签订|$)",
        r"出卖人(?:\(乙方\)|（乙方）)?[：:]\s*(.+?)(?:\s|签订|买受人|$)",
        r"乙方(?:\(出卖方\)|（出卖方）)?[：:]\s*(.+?)(?:\s|电话|地址|$)",
    ]:
        m = re.search(pattern, text)
        if m:
            info["seller"] = m.group(1).strip()
            break

    # 表格内容
    if tables:
        table = tables[0]
        headers = [h.strip() if h else "" for h in table[0]]
        header_map = {h: i for i, h in enumerate(headers)}

        def get(row, col_list):
            for c in col_list:
                if c in header_map:
                    val = row[header_map[c]]
                    return val.strip() if val else ""
            return ""

        for row in table[1:]:
            if not any(cell and str(cell).strip() for cell in row):
                continue
            # 跳过合计/大写行
            row_text = "".join(str(c) or "" for c in row)
            if "大写" in row_text or "合计" in row_text or "总计" in row_text:
                continue

            product = get(row, ["物料名称", "设备名称", "产品名称", "货物名称", "品名", "名称"])
            if not product:
                continue
            unit_price_str = get(row, ["含税单价", "单价"])
            total_price_str = get(row, ["含税金额", "金额", "总价", "合计"])
            if not unit_price_str:
                unit_price_str = get(row, ["无税单价"])
            if not total_price_str:
                total_price_str = get(row, ["无税金额"])
            try:
                unit_price = float(re.sub(r"[^\d.]", "", unit_price_str)) if unit_price_str else 0.0
                total_price = float(re.sub(r"[^\d.]", "", total_price_str)) if total_price_str else 0.0
            except ValueError:
                continue

            info["items"].append({
                "product": product,
                "model": get(row, ["型号规格", "型号"]),
                "unit": get(row, ["单位"]),
                "quantity": get(row, ["数量"]),
                "unit_price": unit_price,
                "total_price": total_price,
            })

        # 备注
        for row in table[1:]:
            row_text = "".join(str(c) or "" for c in row)
            if "备注" in row_text:
                info["note"] = row_text.replace("备注", "").replace("：", "").strip()
                break

    return info, text


# ---------------------------------------------------------------------------
# 图片 OCR 信息提取（EasyOCR）
# ---------------------------------------------------------------------------
def extract_image_info(image_path: str):
    from PIL import Image
    import numpy as np
    import easyocr

    img = Image.open(image_path)
    arr = np.array(img)
    reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
    raw = reader.readtext(arr, detail=1)

    # 按行（y 坐标）聚类
    lines = []
    for bbox, text, conf in raw:
        if not text.strip():
            continue
        cy = (bbox[0][1] + bbox[2][1]) / 2
        cx = (bbox[0][0] + bbox[2][0]) / 2
        lines.append((cy, cx, text.strip()))

    # 简单按 y 分组：相邻 cy 差 < 20 视为同一行
    lines.sort(key=lambda x: x[0])
    grouped = []
    for cy, cx, text in lines:
        if not grouped or abs(cy - grouped[-1]["y"]) > 20:
            grouped.append({"y": cy, "cells": [(cx, text)]})
        else:
            grouped[-1]["cells"].append((cx, text))
    # 每行按 x 排序并拼接
    row_texts = []
    for g in grouped:
        g["cells"].sort(key=lambda x: x[0])
        row_texts.append(" ".join(t for _, t in g["cells"]))
    full_text = "\n".join(row_texts)

    info = {
        "buyer": "",
        "seller": "",
        "order_date": parse_date(full_text),
        "items": [],
        "note": "",
    }

    # 买方/卖方
    for pattern in [
        r"买受方[：:]\s*(.+?)(?:\s|签订|$)",
        r"买受人[：:]\s*(.+?)(?:\s|签订|$)",
        r"甲方[：:]\s*(.+?)(?:\s+乙方|$)",
    ]:
        m = re.search(pattern, full_text)
        if m:
            info["buyer"] = m.group(1).strip()
            break

    for pattern in [
        r"出卖方[：:]\s*(.+?)(?:\s|签订|$)",
        r"出卖人[：:]\s*(.+?)(?:\s|签订|$)",
        r"乙方[：:]\s*(.+?)(?:\s|电话|地址|$)",
    ]:
        m = re.search(pattern, full_text)
        if m:
            info["seller"] = m.group(1).strip()
            break

    # 表格行解析：尝试识别“产品名 + 价格 + 价格”模式
    # 按行逐个检查：如果一行包含数字价格（≥2 个），且前面有非数字文本，视为数据行
    for row in row_texts:
        nums = re.findall(r"\b\d{2,}(?:\.\d{1,2})?\b", row)
        if len(nums) >= 2:
            # 取前两个非数字片段作为产品名，其余作为型号等
            parts = re.split(r"\s+", row.strip())
            # 过滤纯价格项
            non_price = [p for p in parts if not re.fullmatch(r"\d{2,}(?:\.\d{1,2})?", p.replace(",", "").replace(".", ""))]
            prices = [float(p.replace(",", "")) for p in parts if re.fullmatch(r"\d{2,}(?:\.\d{1,2})?", p.replace(",", ""))]
            if len(prices) < 2:
                continue
            # 产品名：取第一个非价格词组
            product = non_price[0] if non_price else ""
            if not product or product in ["头轮总成", "尾轮总成", "涨紧轮总成"] == False:
                # 若无法识别，跳过
                pass
            # 取最后两个价格分别作为单价和总价
            unit_price = prices[-2]
            total_price = prices[-1]
            # 简单型号提取：找 THYxx-xxxx 模式
            m = re.search(r"THY\d{2}[-–—]\d{4}", row, re.IGNORECASE)
            model = m.group(0) if m else ""
            info["items"].append({
                "product": product,
                "model": model,
                "unit": "套",
                "quantity": "1",
                "unit_price": unit_price,
                "total_price": total_price,
            })

    return info, full_text


# ---------------------------------------------------------------------------
# 统一入口
# ---------------------------------------------------------------------------
def extract_file_info(file_path: str):
    ext = Path(file_path).suffix.lower()
    if ext in [".jpg", ".jpeg", ".png", ".bmp"]:
        return extract_image_info(file_path)
    return extract_pdf_info(file_path)


# ---------------------------------------------------------------------------
# Word 模板编辑（支持多行物料）
# ---------------------------------------------------------------------------
def _clone_row_before(table, src_tr, target_tr):
    """在 target_tr 之前插入 src_tr 的克隆。"""
    new_tr = deepcopy(src_tr)
    target_tr.addprevious(new_tr)


def _remove_row(table, row):
    table._tbl.remove(row._tr)


def edit_contract(docx_path: str, info: dict, sales_no: str, discount: float):
    doc = Document(docx_path)

    order_date = info["order_date"]
    signing_date = order_date + timedelta(days=5)
    delivery_date = order_date + timedelta(days=60)

    items = info["items"]
    if not items:
        raise ValueError("未能从文件1提取到任何产品行")

    # 每项打折
    for it in items:
        it["discounted_unit"] = round(it["unit_price"] * discount, 2)
        it["discounted_total"] = round(it["total_price"] * discount, 2)

    discounted_total = round(sum(it["discounted_total"] for it in items), 2)

    contract_no = f"HPGH{signing_date.strftime('%y%m%d')}-1"
    signing_str = signing_date.strftime("%y/%m/%d")
    delivery_str = f"{delivery_date.year}年{delivery_date.month}月{delivery_date.day}日前具备发货条件"
    total_chinese = money_to_chinese(discounted_total)
    total_decimal_str = f"{discounted_total:.2f}"
    buyer_short = short_buyer(info["buyer"])

    # 1) 合同编号
    for p in doc.paragraphs:
        if "合同编号" in p.text and re.search(r"HPGH\d{6}-\d+", p.text):
            replace_in_paragraph(p, re.search(r"HPGH\d{6}-\d+", p.text).group(0), contract_no)
            break

    # 2) 签订时间
    for p in doc.paragraphs:
        if "签订时间" in p.text and re.search(r"\d{2}/\d{2}/\d{2}", p.text):
            replace_in_paragraph(p, re.search(r"\d{2}/\d{2}/\d{2}", p.text).group(0), signing_str)
            break

    # 3) 第一条表格
    if doc.tables:
        table = doc.tables[0]
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        tr_elements = table._tbl.findall(f".//{W}tr")
        row_objects = [_Row(tr, table) for tr in tr_elements]

        # 定位关键行
        data_row_idx = None
        remark_row_idx = None
        total_row_idx = None
        for i, row in enumerate(row_objects):
            row_text = "".join(c.text for c in row.cells)
            if i == 0:
                continue
            if re.search(r"\b1990000\b", row_text) or "玉米除石设备" in row_text or "THY" in row_text:
                if data_row_idx is None:
                    data_row_idx = i
            if "备注" in row_text and remark_row_idx is None:
                remark_row_idx = i
            if "大写" in row_text and total_row_idx is None:
                total_row_idx = i

        if data_row_idx is None and len(row_objects) > 1:
            data_row_idx = 1
        if total_row_idx is None:
            total_row_idx = len(row_objects) - 1

        # 多行时删除备注行以腾出空间
        if len(items) > 1 and remark_row_idx is not None:
            table._tbl.remove(tr_elements[remark_row_idx])
            # 重新获取 tr 列表
            tr_elements = table._tbl.findall(f".//{W}tr")
            row_objects = [_Row(tr, table) for tr in tr_elements]
            if total_row_idx > remark_row_idx:
                total_row_idx -= 1

        # 如果数据行不足，在总计行前克隆
        needed = len(items)
        existing_data_rows = total_row_idx - data_row_idx
        for _ in range(needed - existing_data_rows):
            src_tr = tr_elements[data_row_idx]
            target_tr = tr_elements[total_row_idx]
            _clone_row_before(table, src_tr, target_tr)
            # 重新获取 tr 列表
            tr_elements = table._tbl.findall(f".//{W}tr")
            row_objects = [_Row(tr, table) for tr in tr_elements]
            total_row_idx += 1

        # 填充数据行
        for idx, it in enumerate(items):
            row = row_objects[data_row_idx + idx]
            set_cell_text(row.cells[0], it["product"])
            set_cell_text(row.cells[1], it["model"] if it["model"] else "/")
            set_cell_text(row.cells[2], it["unit"])
            set_cell_text(row.cells[3], it["quantity"])
            set_cell_text(row.cells[4], str(int(it["discounted_unit"])))
            set_cell_text(row.cells[5], str(int(it["discounted_total"])))
            set_cell_text(row.cells[6], buyer_short)
            set_cell_text(row.cells[7], delivery_str)

        # 填充大写金额行
        total_row = row_objects[total_row_idx]
        for cell in total_row.cells:
            if "大写" in cell.text:
                new_text = f"大写：{total_chinese}。（含13%增值税） 总计：（小写）{total_decimal_str}元；"
                set_cell_text(cell, new_text)
                break

    # 4) 结算条款：替换 BGJ 开头的销售合同号
    for p in doc.paragraphs:
        m = re.search(r"BGJ[^\\s]*", p.text)
        if m:
            replace_in_paragraph(p, m.group(0), sales_no)
            break

    # 5) 紧凑排版：尽量保持一页
    _compact_layout(doc)

    return doc, discounted_total


def _compact_layout(doc):
    """调小表格行高、正文行距和页边距，使多行物料也能在一页内显示。"""
    # 页边距：保持原模板 1.27cm，仅底部略微收紧以挤出空间
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # 标题字号略收
    for p in doc.paragraphs:
        if p.text.strip().replace(" ", "") == "工矿机械产品买卖合同":
            for r in p.runs:
                r.font.size = Pt(16)
            p.paragraph_format.space_after = Pt(2)
            break

    if doc.tables:
        # 产品表
        table = doc.tables[0]
        for row in table.rows:
            row.height = Cm(0.48)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(9)

        # 签章信息表
        if len(doc.tables) > 1:
            for row in doc.tables[1].rows:
                row.height = Cm(2.0)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        p.paragraph_format.line_spacing = Pt(15)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                        for r in p.runs:
                            if r.font.size is None:
                                r.font.size = Pt(10.5)

    # 正文固定行距 15 磅（标题、表格外段落）
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt or txt.replace(" ", "") == "工矿机械产品买卖合同":
            continue
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(15)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------------------
# Word COM 保存 .doc / 桌面副本
# ---------------------------------------------------------------------------
def save_as_doc(docx_path: str, doc_path: str):
    import win32com.client as win32

    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(Path(docx_path).resolve()), False, True)
        doc.SaveAs(str(Path(doc_path).resolve()), 0)  # wdFormatDocument
        doc.Close(False)
    finally:
        word.Quit()


def copy_to_desktop(src: str):
    desktop = Path("C:/Users/yesli/OneDrive/桌面")
    desktop.mkdir(parents=True, exist_ok=True)
    dest = desktop / Path(src).name
    shutil.copy2(src, dest)
    return str(dest)


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="采购订单/销售合同 → 供应商采购合同")
    parser.add_argument("file1", help="采购订单/销售合同 PDF 或图片路径")
    parser.add_argument("template", help="Word 合同模板路径（.doc/.docx）")
    parser.add_argument("--output-dir", "-o", default=".", help="输出目录，默认当前目录")
    parser.add_argument("--discount", "-d", type=float, default=0.8, help="折扣，默认 0.8")
    parser.add_argument("--buyer", default="", help="手动指定买方全称（OCR 不准时使用）")
    parser.add_argument("--seller", default="", help="手动指定卖方全称（OCR 不准时使用）")
    parser.add_argument("--order-date", default="", help="手动指定签订/订单日期 YYYY-MM-DD")
    parser.add_argument("--items-json", default="", help="手动指定物料 JSON 列表（覆盖 OCR）")
    args = parser.parse_args()

    file1 = Path(args.file1)
    template = Path(args.template)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if not file1.exists():
        print(f"错误：找不到文件1 {file1}", file=sys.stderr)
        sys.exit(1)
    if not template.exists():
        print(f"错误：找不到模板 {template}", file=sys.stderr)
        sys.exit(1)

    sales_no = extract_sales_contract_no(file1.name)

    # 如果提供了完整的物料和日期，跳过文件内容提取（避免依赖 OCR）
    if args.items_json and args.buyer and args.order_date:
        info = {
            "buyer": args.buyer,
            "seller": args.seller or "",
            "order_date": datetime.strptime(args.order_date, "%Y-%m-%d").date(),
            "items": json.loads(args.items_json),
            "note": "",
        }
    else:
        info, _ = extract_file_info(str(file1))
        if args.buyer:
            info["buyer"] = args.buyer
        if args.seller:
            info["seller"] = args.seller
        if args.order_date:
            info["order_date"] = datetime.strptime(args.order_date, "%Y-%m-%d").date()
        if args.items_json:
            info["items"] = json.loads(args.items_json)

    if not info["buyer"]:
        print("警告：未能提取买方名称，可用 --buyer 指定", file=sys.stderr)
    if info["order_date"] is None:
        print("错误：未能提取签订/订单日期，可用 --order-date 指定", file=sys.stderr)
        sys.exit(1)
    if not info["items"]:
        print("错误：未能提取产品行，可用 --items-json 指定", file=sys.stderr)
        sys.exit(1)

    buyer_for_filename = re.sub(r'[\\/:*?"<>|]', '', info["buyer"]).strip() or short_buyer(info["buyer"])
    discounted_total = round(sum(it["total_price"] for it in info["items"]) * args.discount, 2)
    today_str = datetime.now().strftime("%y%m%d")
    out_filename = f"{today_str}-hpgh-{buyer_for_filename}-{int(discounted_total)}.doc"
    out_path = output_dir / out_filename

    # 临时目录
    tmp_dir = output_dir / f"_tmp_contract_{datetime.now().strftime('%H%M%S')}"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    tmp_docx = tmp_dir / "template.docx"

    # Word COM 转 .doc → .docx
    import win32com.client as win32

    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(template.resolve()), False, True)
        doc.SaveAs(str(tmp_docx.resolve()), 16)
        doc.Close(False)
    finally:
        word.Quit()

    # 编辑
    edited_docx = tmp_dir / "edited.docx"
    doc, _ = edit_contract(str(tmp_docx), info, sales_no, args.discount)
    doc.save(str(edited_docx))

    # 另存为 .doc
    save_as_doc(str(edited_docx), str(out_path))

    # 桌面副本
    desktop_path = copy_to_desktop(str(out_path))

    # 清理临时文件
    shutil.rmtree(tmp_dir, ignore_errors=True)

    print(f"生成合同：{out_path}")
    print(f"桌面副本：{desktop_path}")
    print(f"销售合同号：{sales_no}")


if __name__ == "__main__":
    main()
