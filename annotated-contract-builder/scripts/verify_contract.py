# -*- coding: utf-8 -*-
"""
合同成品程序化校验：
  - 必备字段全部存在（--require）
  - 旧值/残留字段全部不存在（--forbid，如旧单位名、旧金额、旧账号、旧日期）
  - 三色标注统计（灰/蓝/绿 run 数）
  - 节数、标题数、图片数

用法：
  python verify_contract.py 合同.docx --require 新单位名 240000.00 贰拾肆万元整 \
      --forbid 旧单位名 158000.00 壹拾伍万
也可用 --config checks.json：{"require": [...], "forbid": [...]}

校验不通过时退出码为 1，并列出每个残留/缺失项。
"""
import argparse
import json
import re
import sys

import docx

TRAILING_PUNCT = "。，、：；,.;:;!！?？"
CN_NUMS = "一二三四五六七八九十"


def check_no_legacy_subitem_numbers(doc):
    """检查正文是否还残留旧式子项编号（(1)(2)(3) / ①②③ / (a)(b)）。
    建设工程合同已统一为点分式（X.Y.Z.N），旧式编号应全部清除。
    只检查以这些符号开头的段落（正文中途出现的不算）。"""
    import re
    issues = []
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("Heading"):
            continue
        t = p.text.strip()
        if re.match(r"^[（(][\da-zA-Z一-四][）)]", t):
            issues.append(f"旧式括号子项编号: {t[:30]}…")
        elif re.match(r"^[①②③④⑤⑥⑦⑧⑨⑩]", t):
            issues.append(f"圈码子项编号: {t[:30]}…")
    return issues


def check_no_source_numbers(doc):
    """结构化检查：正文条款前缀 X.Y / X.Y.Z 的 X 必须与最近的「第X条」标题匹配。
    这能发现来源编号混入（如第四条下出现 6.2.x），同时不误报合同统一编号。"""
    import re
    issues = []
    current_tiao = None
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        t = p.text.strip()
        if st == "Heading 1":
            m = re.match(r"^第([一二三四五六七八九十百]+)条", t)
            if m:
                cn_map = {"一":1,"二":2,"三":3,"四":4,"五":5,"六":6,
                          "七":7,"八":8,"九":9,"十":10}
                s = m.group(1)
                current_tiao = cn_map.get(s, None)
            continue
        # 检查正文条款前缀：X.Y 或 X.Y.Z
        m = re.match(r"^(\d+)\.(\d+)", t)
        if m and current_tiao:
            clause_major = int(m.group(1))
            if clause_major != current_tiao:
                issues.append(
                    f"编号不匹配: 条款{m.group(0)}不在第{current_tiao}条下"
                    f"（当前为第{current_tiao}条）: {t[:30]}…")
    return issues


def check_headings(doc):
    """标题规范检查（交付前必过）：
    - 所有标题末尾不得带标点（。：，、等）
    - 一级标题（第X条、XXX）标题部分汉字数 ≤ 12
    - 一级标题按 第一条..第N条 连续编号（同一套体系，不混来源编号）
    """
    issues = []
    h1s = []
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        txt = p.text.strip()
        if st not in ("Heading 1", "Heading 2", "Heading 3") or not txt:
            continue
        if txt[-1] in TRAILING_PUNCT:
            issues.append(f"标题末尾带标点: {txt[:30]}…" if len(txt) > 30
                          else f"标题末尾带标点: {txt}")
        if st == "Heading 1":
            h1s.append(txt)
            m = re.match(r"^第[一二三四五六七八九十百]+条、?", txt)
            title = txt[m.end():] if m else txt
            n_han = len(re.findall(r"[\u4e00-\u9fff]", title))
            if n_han > 12:
                issues.append(f"一级标题超12个汉字({n_han}字): {txt}")
    for i, h in enumerate(h1s, 1):
        if i <= 10:
            expect = f"第{CN_NUMS[i - 1]}条"
        else:
            expect = f"第十{CN_NUMS[i - 11]}条"
        if not h.startswith(expect):
            issues.append(f"一级标题顺序异常: 第{i}个应为「{expect}…」，实际「{h[:20]}…」")
    return issues


def collect(doc):
    texts, colors, headings = [], {}, {1: 0, 2: 0, 3: 0}
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        for lv in (1, 2, 3):
            if st == f"Heading {lv}":
                headings[lv] += 1
        for r in p.runs:
            c = str(r.font.color.rgb) if r.font.color and r.font.color.rgb else "None"
            colors[c] = colors.get(c, 0) + 1
            texts.append(r.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        c = str(r.font.color.rgb) if r.font.color and r.font.color.rgb else "None"
                        colors[c] = colors.get(c, 0) + 1
                        texts.append(r.text)
    return "".join(texts), colors, headings


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--require", nargs="*", default=[])
    ap.add_argument("--forbid", nargs="*", default=[])
    ap.add_argument("--config", default=None)
    args = ap.parse_args()

    require, forbid = list(args.require), list(args.forbid)
    if args.config:
        with open(args.config, encoding="utf-8") as f:
            cfg = json.load(f)
        require += cfg.get("require", [])
        forbid += cfg.get("forbid", [])

    doc = docx.Document(args.docx)
    text, colors, headings = collect(doc)

    print("sections:", len(doc.sections))
    print("headings H1/H2/H3:", headings)
    print("images:", len(doc.inline_shapes))
    print("colors 灰808080/蓝0000FF/绿008000:",
          {k: colors.get(k, 0) for k in ("808080", "0000FF", "008000")})

    ok = True
    issues = check_headings(doc)
    for msg in issues:
        print("FAIL heading:", msg)
        ok = False
    if not issues:
        print("OK   headings: 末尾无标点 / 一级≤12汉字 / 第X条顺序连续")
    src_issues = check_no_source_numbers(doc)
    for msg in src_issues:
        print("FAIL source#:", msg)
        ok = False
    if not src_issues:
        print("OK   条款编号与所属第X条匹配，无来源编号混入")
    legacy = check_no_legacy_subitem_numbers(doc)
    for msg in legacy:
        print("FAIL legacy#:", msg)
        ok = False
    if not legacy:
        print("OK   无旧式子项编号((1)/①/(a))，全点分式")
    for k in require:
        n = text.count(k)
        print(("OK  " if n else "MISS") + f" require {k!r}: {n}")
        if not n:
            ok = False
    for b in forbid:
        n = text.count(b)
        if n:
            print(f"FAIL forbid {b!r}: 残留 {n} 处")
            ok = False
    print("RESULT:", "PASS" if ok else "FAIL")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
