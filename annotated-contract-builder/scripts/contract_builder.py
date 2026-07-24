# -*- coding: utf-8 -*-
"""
三色标注合同生成引擎（annotated-contract-builder）

把"标注颜色 + 国标标题体系 + 封面/目录/页码"的排版机制沉淀为可复用类，
每个具体合同只需编写一个内容脚本调用 ContractBuilder 的 API 即可。

标注颜色约定（由 SKILL.md 与 references/annotation-rules.md 定义）：
  GRAY  灰 = 模板未改动内容，或在模板基础上改写的内容（含金额、日期、地点等指定改动）
  BLUE  蓝 = 模板内容被说明材料详细内容替换的部分（尽量保持原文）
  GREEN 绿 = 在说明材料内容基础上进一步改写/补全/纠错/适配的部分

版面结构：
  封面节（无页码）→ 目录节（无页码，TOC 域）→ 正文节（页脚居中页码，从 1 起编）
  cover() 与 toc() 均可选；不用封面时直接在首节放目录，不用目录时正文从首节开始。

依赖：python-docx。媒体提取用标准库 zipfile。
"""
import os
import zipfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GRAY = RGBColor(0x80, 0x80, 0x80)
BLUE = RGBColor(0x00, 0x00, 0xFF)
GREEN = RGBColor(0x00, 0x80, 0x00)

BODY_FONT = "宋体"
HEAD_FONT = "黑体"


class ContractBuilder:
    """三色标注合同构建器。用法见 examples/ 下的内容脚本。"""

    def __init__(self):
        self.doc = Document()
        sec = self.doc.sections[0]
        # A4 + 常规合同页边距
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(2.54)
        sec.left_margin = sec.right_margin = Cm(3.0)
        self.body_section = None

    # ---------- 基础构件 ----------

    def _style_run(self, r, color, size, font_cn=BODY_FONT, bold=False):
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = "Times New Roman"
        if color is not None:
            r.font.color.rgb = color
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)

    def para(self, runs, indent=True, bold=False, align=None, size=12,
             space_after=0, space_before=0, line_spacing=1.5):
        """正文段落。runs 为 (text, color) 或 [(text, color), ...]，
        支持同一段内多色混排（蓝原文 + 绿改写词）。"""
        if isinstance(runs, tuple):
            runs = [runs]
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        if indent:
            pf.first_line_indent = Pt(size * 2)  # 首行缩进2字符
        pf.space_after = Pt(space_after)
        pf.space_before = Pt(space_before)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
        if align is not None:
            p.alignment = align
        for text, color in runs:
            r = p.add_run(text)
            self._style_run(r, color, size, bold=bold)
        return p

    def h(self, level, text, color):
        """国标标题：1级=第X条(黑体四号)，2级=X.Y(黑体小四)，3级=X.Y.Z(宋体小四加粗)。
        使用 Word 内置 Heading 样式以获得目录大纲级别；颜色保留标注色。"""
        size = {1: 14, 2: 12, 3: 12}[level]
        font_cn = HEAD_FONT if level <= 2 else BODY_FONT
        bold = False if level <= 2 else True
        p = self.doc.add_paragraph(style=f"Heading {level}")
        pf = p.paragraph_format
        pf.space_before = Pt(12 if level == 1 else 6)
        pf.space_after = Pt(6 if level == 1 else 3)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.first_line_indent = Pt(0)
        r = p.add_run(text)
        self._style_run(r, color, size, font_cn=font_cn, bold=bold)
        return p

    def h1(self, text, color=GRAY):
        return self.h(1, text, color)

    def h2(self, text, color=GRAY):
        return self.h(2, text, color)

    def h3(self, text, color=GRAY):
        return self.h(3, text, color)

    def clause(self, number, runs, indent=True, size=12,
               space_after=0, line_spacing=1.5):
        """带统一编号的正文条款：在段落最前面加 'number ' 前缀。
        用于把来源编号（6.2.1、7.1 等）替换为合同体系统一编号（4.1、5.1 等）。
        runs 用法同 para()。"""
        if isinstance(runs, tuple):
            runs = [runs]
        first_text, first_color = runs[0]
        runs[0] = (f"{number} {first_text}", first_color)
        return self.para(runs, indent=indent, size=size,
                         space_after=space_after, line_spacing=line_spacing)

    def subitem(self, number, text, color, size=12):
        """条款下的子项，用点分编号（如 6.4.1.1、8.2.1），与 clause 的 X.Y.Z 体系一致。"""
        return self.para((f"{number} {text}", color), indent=True, size=size)

    def image(self, path, width_cm=10.5):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run()
        r.add_picture(path, width=Cm(width_cm))

    def blank(self, n=1):
        for _ in range(n):
            self.doc.add_paragraph()

    # ---------- 标注图例 ----------

    def legend(self, source_name="说明材料", size=10.5):
        """在正文开头插入三色标注说明。source_name 为说明材料名称（如《XX安装说明》）。"""
        self.para([("标注说明：", GRAY),
                   ("灰色＝模板未改动或在模板基础上改写的内容；", GRAY),
                   (f"蓝色＝由{source_name}替换的详细内容；", BLUE),
                   ("绿色＝在说明材料内容基础上改写的内容。", GREEN)],
                  indent=False, size=size)

    # ---------- 封面 / 目录 / 页码 ----------

    def cover(self, title, lines, title_size=22, top_space=150, line_size=14,
              title_color=GRAY):
        """封面（当前节），随后开新节。lines 为居中信息行列表，如
        ['合同编号：HP 20260720', '甲方（发包单位）：XX公司', '乙方（承包方）：', ...]"""
        p = self.para((title, title_color), indent=False, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, size=title_size)
        p.paragraph_format.space_before = Pt(top_space)
        for i, line in enumerate(lines):
            q = self.para((line, GRAY), indent=False,
                          align=WD_ALIGN_PARAGRAPH.CENTER, size=line_size,
                          space_after=6)
            if i == 0:
                q.paragraph_format.space_before = Pt(80)
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    def _add_field(self, paragraph, instr_text, placeholder=None,
                   color=None, size=10.5):
        r = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instr_text
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r._r.append(fld_begin)
        r._r.append(instr)
        r._r.append(fld_sep)
        if placeholder is not None:
            r2 = paragraph.add_run(placeholder)
            self._style_run(r2, color, size)
        r3 = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r3._r.append(fld_end)

    def toc(self, levels="1-1", title="目　　录", title_color=GRAY,
            compact=True):
        """目录页（当前节）：标题 + TOC 域，随后开正文节并设置页码从 1 起编。
        TOC 域页码需用 scripts/update_fields.py（Word COM）刷新后显示。
        levels 默认 "1-1"（目录只显示一级标题）。
        compact=True 时定义紧凑的 TOC 1/2/3 样式（小四→五号、单倍行距、段距0），
        确保目录压缩在一页以内；条目更多时可再把字号调小。"""
        self.para((title, title_color), indent=False, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=16, space_after=12)
        if compact:
            self._define_toc_styles()
        p = self.doc.add_paragraph()
        self._add_field(p, f'TOC \\o "{levels}" \\h \\z \\u',
                        placeholder="（目录将在文档打开后自动生成）")
        self.body_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self._setup_body_page_numbers(self.body_section)

    def _define_toc_styles(self, font_size=10.5):
        """定义紧凑 TOC 样式。Word 刷新目录域时按样式名（toc 1/2/3）套用，
        以此控制目录字号与行距，把目录压到一页内。"""
        from docx.enum.style import WD_STYLE_TYPE
        for lv in (1, 2, 3):
            name = f"TOC {lv}"
            try:
                st = self.doc.styles[name]
            except KeyError:
                st = self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.font.size = Pt(font_size)
            st.font.name = "Times New Roman"
            rpr = st.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), BODY_FONT)
            pf = st.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(2)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.0

    def _setup_body_page_numbers(self, section):
        """正文节：页脚居中 PAGE 域，页码从 1 重新起编。"""
        sectPr = section._sectPr
        pgNumType = OxmlElement("w:pgNumType")
        pgNumType.set(qn("w:start"), "1")
        sectPr.append(pgNumType)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_field(p, "PAGE", placeholder="1", color=GRAY)

    # ---------- 签署栏 ----------

    def signature_table(self, rows, size=12):
        """两行签署栏：rows 为 [(左单元格文本, 右单元格文本), ...]。
        留空方信息传空字符串或只留标签（如 '乙方：'）。"""
        self.blank(1)
        table = self.doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (l, r_) in enumerate(rows):
            for j, txt in enumerate((l, r_)):
                cell = table.cell(i, j)
                p = cell.paragraphs[0]
                run = p.add_run(txt)
                self._style_run(run, GRAY, size)
        return table

    def save(self, path):
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        self.doc.save(path)
        return path


# ---------- 工具函数 ----------

def extract_media(docx_path, out_dir):
    """从说明材料 docx 提取嵌入图片到 out_dir，返回文件名列表（按名称排序）。
    用于把说明材料中的设备图/示意图原样嵌入新合同。"""
    os.makedirs(out_dir, exist_ok=True)
    names = []
    with zipfile.ZipFile(docx_path) as z:
        for n in z.namelist():
            if n.startswith("word/media/"):
                data = z.read(n)
                fn = os.path.basename(n)
                with open(os.path.join(out_dir, fn), "wb") as f:
                    f.write(data)
                names.append(fn)
    return sorted(names)


def image_anchors(docx_path):
    """定位说明材料中每张图片锚定的段落位置。
    返回 [(段落索引, 段落前文文本, [rId...]), ...] 及 rId→文件名映射，
    用于判断'如下图'引用的是哪张图。"""
    import docx
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn as _qn
    d = docx.Document(docx_path)
    anchors = []
    prev_text = ""
    for i, child in enumerate(d.element.body.iterchildren()):
        if child.tag == _qn("w:p"):
            p = Paragraph(child, d)
            blips = child.findall(".//" + _qn("a:blip"))
            if blips:
                rids = [b.get(_qn("r:embed")) for b in blips]
                anchors.append((i, prev_text, rids))
            elif p.text.strip():
                prev_text = p.text.strip()
    rels = {rid: rel.target_ref for rid, rel in d.part.rels.items()
            if "image" in rel.reltype}
    return anchors, rels
