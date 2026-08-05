# -*- coding: utf-8 -*-
"""
根据模板和配置 JSON 生成国华销售合同。
"""
import json
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, size=10.5, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def replace_paragraph_text(p, new_text, size=10.5, bold=False):
    """清空段落内所有 runs，替换为单一段落文本。"""
    for r in p.runs:
        r.text = ""
    if not p.runs:
        p.add_run()
    p.runs[0].text = new_text
    set_run_font(p.runs[0], size, bold)


def center_cell_text(cell, size=10.5):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_run_font(r, size)


def update_sign_cell(cell, lines, size=10.5):
    """
    在签署栏单元格的“电话”段落后插入“传真”行，并更新所有行文本。
    假设原单元格至少有 9 段：标题、名称、住所、代理人、电话、银行、帐号、税号、空行。
    """
    paras = list(cell.paragraphs)

    # 先修改前 5 段
    for i in range(min(5, len(paras))):
        replace_paragraph_text(paras[i], lines[i], size)

    # 在第 4 段（电话）后插入传真段落
    if len(paras) > 4:
        phone_p = paras[4]._p
        new_p_xml = OxmlElement("w:p")
        phone_p.addnext(new_p_xml)
        from docx.text.paragraph import Paragraph

        fax_p = Paragraph(new_p_xml, cell)
        replace_paragraph_text(fax_p, lines[5], size)

    # 刷新 paragraphs 列表，修改后续段落
    paras = list(cell.paragraphs)
    for idx, line_idx in enumerate([6, 7, 8], start=6):
        if idx < len(paras):
            replace_paragraph_text(paras[idx], lines[line_idx], size)

    # 删除末尾空段落
    while cell.paragraphs and cell.paragraphs[-1].text.strip() == "":
        last_p = cell.paragraphs[-1]._p
        last_p.getparent().remove(last_p)


def update_table_items(table, items, total_text):
    """
    更新标的表格。
    假设表格结构：表头 + 数据行 + 金额行。
    金额行为合并单元格：第一列"人民币"，第二列为金额文本。
    """
    # 简单假设：最后行为金额行，前面是数据行
    data_rows = len(table.rows) - 1

    # 替换现有数据行
    for i, item in enumerate(items):
        if i + 1 >= data_rows:
            break
        row = table.rows[i + 1].cells
        vals = [item.get(k, "") for k in ("name", "spec", "unit", "qty", "price", "amount", "note")]
        # 如果表格没有"备注"列，去掉最后一个
        if len(row) < 7:
            vals = vals[: len(row)]
        for col_idx, v in enumerate(vals):
            row[col_idx].text = str(v)
            center_cell_text(row[col_idx])

    # 清空多余数据行
    for i in range(len(items) + 1, data_rows):
        row = table.rows[i].cells
        for cell in row:
            cell.text = ""

    # 金额行
    amount_row = table.rows[-1].cells
    amount_row[0].text = "人民币"
    amount_row[1].text = total_text
    for i in range(2):
        center_cell_text(amount_row[i])


def build_contract(config_path):
    with open(config_path, "r", encoding="utf-8") as f:
        cfg = json.load(f)

    doc = Document(cfg["template_docx"])

    seller = cfg["seller"]
    buyer = cfg["buyer"]

    # 修改抬头
    for p in doc.paragraphs:
        text = p.text
        if "合同编号：" in text and cfg.get("contract_no"):
            # 保留原格式，只替换编号部分
            new_text = text
            import re

            new_text = re.sub(r"合同编号：\S+", f"合同编号：{cfg['contract_no']}", new_text)
            replace_paragraph_text(p, new_text)
        if "签订时间：" in text and cfg.get("sign_date"):
            new_text = re.sub(r"签订时间：\S+", f"签订时间：{cfg['sign_date']}", text)
            replace_paragraph_text(p, new_text)
        # 如需替换双方名称可在此扩展

    # 替换标的表格
    if cfg.get("items") and cfg.get("total_text"):
        update_table_items(doc.tables[0], cfg["items"], cfg["total_text"])

    # 修改签署栏
    left_lines = [
        "出卖方",
        f"出卖人：{seller.get('name', '')}",
        f"住所：{seller.get('address', '')}",
        f"委托代理人：{seller.get('agent', '')}",
        f"电话：{seller.get('phone', '')}",
        f"传真：{seller.get('fax', '')}",
        f"开户银行：{seller.get('bank', '')}",
        f"帐号：{seller.get('account', '')}",
        f"税号：{seller.get('tax', '')}",
    ]
    right_lines = [
        "买受方",
        f"买受人：{buyer.get('name', '')}",
        f"住所：{buyer.get('address', '')}",
        f"委托代理人：{buyer.get('agent', '')}",
        f"电话：{buyer.get('phone', '')}",
        f"传真：{buyer.get('fax', '')}",
        f"开户银行：{buyer.get('bank', '')}",
        f"帐号：{buyer.get('account', '')}",
        f"税号：{buyer.get('tax', '')}",
    ]

    if len(doc.tables) > 1:
        update_sign_cell(doc.tables[1].rows[0].cells[0], left_lines)
        update_sign_cell(doc.tables[1].rows[0].cells[1], right_lines)

    os.makedirs(os.path.dirname(os.path.abspath(cfg["output_docx"])), exist_ok=True)
    doc.save(cfg["output_docx"])
    print(f"saved: {cfg['output_docx']}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: build_contract.py <config.json>")
        sys.exit(1)
    build_contract(sys.argv[1])
