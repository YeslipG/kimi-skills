from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import re
import os


def set_run_font(run, font_name='宋体', size=10.5, bold=False, underline=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_margins(cell, top=0, left=0, bottom=0, right=0):
    """设置单元格内边距，消除默认内边距浪费的空间"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        mar = OxmlElement(f'w:{edge}')
        mar.set(qn('w:w'), str(val))
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)


def num_to_chinese(num):
    """数字转中文大写金额"""
    num = int(num)
    if num == 0:
        return '零元整'
    units = ['', '拾', '佰', '仟']
    nums = '零壹贰叁肆伍陆柒捌玖'
    result = ''
    str_num = str(num)
    length = len(str_num)
    zero_flag = False
    for i, ch in enumerate(str_num):
        n = int(ch)
        pos = length - i - 1
        if n == 0:
            if not zero_flag and pos % 4 == 0 and pos > 0:
                result += '零'
                zero_flag = True
        else:
            zero_flag = False
            result += nums[n] + units[pos % 4]
    result = result.replace('零零', '零').replace('零万', '万').replace('零亿', '亿')
    result = result.rstrip('零') + '元整'
    return result


def generate_contract(
    buyer_name,
    product_name,
    product_model,
    unit,
    quantity,
    unit_price,
    delivery_time,
    seller_name='南通厚普机械科技有限公司',
    contract_prefix='HP',
    buyer_tax='',
    buyer_address='',
    buyer_phone='',
    buyer_bank='',
    buyer_account=''
):
    """生成工矿机械产品买卖合同（严格1页内）"""
    doc = Document()
    section = doc.sections[0]
    # 页边距：上下2.0cm，左右3.0cm
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    today = datetime.date.today()
    date_str = today.strftime('%Y-%m-%d')
    date_short = today.strftime('%Y%m%d')
    total_price = quantity * unit_price
    contract_no = f"{contract_prefix}{date_short}"

    # 买方简称
    buyer_short = buyer_name
    buyer_short = re.sub(r'(河北省|邢台|南通|启东|北京|上海|广州|深圳|江苏省|浙江|山东|山西|河南|湖北|湖南|广东|四川|贵州|云南|陕西|甘肃|青海|台湾|内蒙古|广西|西藏|宁夏|新疆|香港|澳门|黑龙江|吉林|辽宁|天津|重庆|江西|福建|安徽|海南|市|省|县|区)', '', buyer_short)
    buyer_short = re.sub(r'(有限公司|有限责任公司|股份公司|股份有限公司)', '', buyer_short)
    buyer_short = re.sub(r'[\(\)（）]', '', buyer_short)
    buyer_short = buyer_short.strip()

    filename = f"{date_short}-配件销售-合同-{buyer_short}-{total_price}"

    # ========== 标题 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.space_before = Pt(0)
    run = title.add_run('工矿机械产品买卖合同')
    set_run_font(run, font_name='宋体', size=20, bold=True)

    # ========== 合同信息行（原版位置） ==========
    info_table = doc.add_table(rows=3, cols=2)
    info_table.autofit = False
    info_table.allow_autofit = False

    # 隐藏边框 + 内边距0
    for row in info_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                edge_el = OxmlElement(f'w:{edge}')
                edge_el.set(qn('w:val'), 'none')
                tcBorders.append(edge_el)
            tcPr.append(tcBorders)
            set_cell_margins(cell, 0, 0, 0, 0)

    for cell in info_table.columns[0].cells:
        cell.width = Cm(8)
    for cell in info_table.columns[1].cells:
        cell.width = Cm(8)

    # 固定行高280twips（约0.49cm）
    for row in info_table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '280')
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

    # 第1行：出卖方（左） | 合同编号（右）
    cell = info_table.cell(0, 0)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f'出卖方：{seller_name}')
    set_run_font(run, font_name='宋体', size=10.5)

    cell = info_table.cell(0, 1)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f'合同编号：{contract_no}')
    set_run_font(run, font_name='宋体', size=10.5)

    # 第2行：买受方（左） | 签订地点（右）
    cell = info_table.cell(1, 0)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f'买受方：{buyer_name}')
    set_run_font(run, font_name='宋体', size=10.5)

    cell = info_table.cell(1, 1)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run('签订地点：南通市')
    set_run_font(run, font_name='宋体', size=10.5)

    # 第3行：空（左） | 签订时间（右）
    cell = info_table.cell(2, 0)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    cell = info_table.cell(2, 1)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f'签订时间：{date_str}')
    set_run_font(run, font_name='宋体', size=10.5)

    # ========== 第一条 表格 ==========
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run('第一条')
    set_run_font(run, font_name='宋体', size=10.5, bold=True)
    run = p.add_run('  标的、数量、价款及交（提）货时间')
    set_run_font(run, font_name='宋体', size=10.5)

    table = doc.add_table(rows=3, cols=7)
    table.style = 'Table Grid'
    table.autofit = False
    widths = [Cm(2.2), Cm(2.8), Cm(1.3), Cm(1.3), Cm(2.2), Cm(2.2), Cm(4.5)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # 表格行高0.6cm（340twips），内边距上下0
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '340')
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        for cell in row.cells:
            set_cell_margins(cell, 0, 55, 0, 55)

    headers = ['名称', '型号规格', '单位', '数量', '单价（元）', '总价（元）', '交货时间']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(header)
        set_run_font(run, font_name='宋体', size=10.5)

    data_row = [product_name, product_model, unit, str(quantity), str(unit_price), str(total_price), delivery_time]
    for i, val in enumerate(data_row):
        cell = table.cell(1, i)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(str(val))
        set_run_font(run, font_name='宋体', size=10.5)

    # 合并人民币行
    for i in range(6, 0, -1):
        table.cell(2, 0).merge(table.cell(2, i))
    cell_merged = table.cell(2, 0)
    p = cell_merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    chinese_total = num_to_chinese(total_price)
    run = p.add_run('人民币')
    set_run_font(run, font_name='宋体', size=10.5)
    run = p.add_run(f'  （小写）¥{total_price}元；（大写）{chinese_total}。（含13%增值税及运费）')
    set_run_font(run, font_name='宋体', size=10.5)

    # ========== 条款（行距18pt固定值，段后0，严格1页） ==========
    clauses = [
        ('第二条', '质量标准：', '按国家标准执行。'),
        ('第三条', '出卖人对质量负责的条件及期限：', '质保期为连续正常使用 12 个月。'),
        ('第四条', '包装标准、包装物的供应与回收：', '出卖人按运输要求免费包装，保证设备完好无损运至买受人指定地点，包装物不回收。'),
        ('第五条', '随机的必备品、配件、工具数量及供应办法：', '随机附带产品相关资料。'),
        ('第六条', '标的物所有权自', '交付起转移。'),
        ('第七条', '交（提）货方式、地点：', '出卖人负责托运至买受人指定地点。'),
        ('第八条', '检验标准、方式、地点及期限：', '按国家标准验收。'),
        ('第九条', '成套设备的安装与调试：', '出卖人负责所供设备的指导安装调试。'),
        ('第十条', '结算方式、时间及地点：', '预付全款加工，发货开具全额增值税发票。'),
        ('第十一条', '担保方式：', '无。'),
        ('第十二条', '合同解除的条件：', '双方协商解决。'),
        ('第十三条', '售后服务：', '质保期内如因产品出现质量问题，出卖人应在收到书面通知后 4 小时内响应；'),
        ('第十四条', '合同争议的解决方式：', '本合同在履行过程中发生的争议，由双方当事人协商解决；协商不成的，依法向买受人所在地人民法院起诉。'),
        ('第十五条', '本合同自双方签字盖章起生效，传真件有效，合同一式两份，合同双方各一份。', ''),
        ('第十六条', '技术协议同本合同具备同等的法律效力。', ''),
        ('第十七条', '其它约定事项：', '无。')
    ]

    for clause_num, clause_title, clause_content in clauses:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(f'{clause_num}  ')
        set_run_font(run, font_name='宋体', size=10.5, bold=True)
        run = p.add_run(clause_title)
        set_run_font(run, font_name='宋体', size=10.5)
        if clause_content:
            run = p.add_run(clause_content)
            set_run_font(run, font_name='宋体', size=10.5, underline=True)

    # ========== 双方信息表格 ==========
    info_table2 = doc.add_table(rows=7, cols=2)
    info_table2.style = 'Table Grid'
    info_table2.autofit = False
    for cell in info_table2.columns[0].cells:
        cell.width = Cm(8)
    for cell in info_table2.columns[1].cells:
        cell.width = Cm(8)

    # 行高0.6cm，内边距上下0
    for row in info_table2.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '340')
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        for cell in row.cells:
            set_cell_margins(cell, 0, 55, 0, 55)

    # 卖方信息数据库
    seller_db = {
        '南通厚普机械科技有限公司': [
            '出卖方', '出卖人：南通厚普机械科技有限公司', '住所：启东市王鲍镇新生街295号',
            '法定代表人：陆锦香', '开户银行：中国工商银行启东市支行',
            '账号：1111 6299 0910 0576 392', '税号：9132 0681 MA1U T2C494'
        ],
        '北京国华技术有限公司': [
            '出卖方', '出卖人：北京国华技术有限公司', '住所：北京市顺义区马坡镇龙跃街9号院1幢6层618室',
            '法定代表人：', '开户银行：中国建设银行股份有限公司北京科创支行',
            '账号：1100 1085 9000 5300 3973', '税号：91110108560429155H'
        ]
    }
    seller_info = seller_db.get(seller_name, [
        '出卖方', f'出卖人：{seller_name}', '住所：', '法定代表人：', '开户银行：', '账号：', '税号：'
    ])

    # 买方信息
    buyer_info = [
        '买受方',
        f'买受人：{buyer_name}',
        f'住所：{buyer_address}' if buyer_address else '住所：',
        f'电话：{buyer_phone}' if buyer_phone else '电话：',
        f'开户银行：{buyer_bank}' if buyer_bank else '开户银行：',
        f'账号：{buyer_account}' if buyer_account else '账号：',
        f'税号：{buyer_tax}' if buyer_tax else '税号：'
    ]

    for i, (s_info, b_info) in enumerate(zip(seller_info, buyer_info)):
        cell = info_table2.cell(i, 0)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(s_info)
            set_run_font(run, font_name='宋体', size=10.5, bold=True)
        else:
            run = p.add_run(s_info)
            set_run_font(run, font_name='宋体', size=10.5)

        cell = info_table2.cell(i, 1)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(b_info)
            set_run_font(run, font_name='宋体', size=10.5, bold=True)
        else:
            run = p.add_run(b_info)
            set_run_font(run, font_name='宋体', size=10.5)

    filepath = f'{filename}.docx'
    doc.save(filepath)
    return filepath, filename, contract_no, total_price


# ========== 执行入口 ==========
if __name__ == '__main__':
    filepath, filename, contract_no, total = generate_contract(
        buyer_name='斯塔克传动技术（天津）有限公司',
        product_name='气囊',
        product_model='QN9649004',
        unit='个',
        quantity=8,
        unit_price=900,
        delivery_time='30天具备发货条件',
        buyer_tax='911201025783048440',
        buyer_address='天津市河东区大桥道52号第A座第二层第A225单元',
        buyer_phone='13389901520',
        buyer_bank='中国银行天津河西支行',
        buyer_account='272664079711'
    )
    print(f'合同已生成: {filepath}')
