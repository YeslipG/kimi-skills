"""
Format an existing Word document according to Chinese technical manual conventions.

This script performs a conservative automatic formatting pass. It is designed to
handle common cases reliably; after running it you should inspect the output and
manually correct any ambiguous paragraphs (especially mixed headings/steps, tables,
headers/footers, and cover pages).

Usage:
    python format_docx.py <input.docx> <output.docx>
"""
import argparse
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Formatting configuration
# ---------------------------------------------------------------------------

HEADING_CN_FONT = '黑体'
BODY_CN_FONT = '宋体'
WESTERN_FONT = 'Times New Roman'
FONT_COLOR = RGBColor(0, 0, 0)

HEADING_SIZES = {
    'Heading 1': Pt(16),   # 三号
    'Heading 2': Pt(15),   # 小三
    'Heading 3': Pt(14),   # 四号
    'Heading 4': Pt(14),   # 四号
}

HEADING_SPACING = {
    'Heading 1': {'before': Pt(18), 'after': Pt(8)},
    'Heading 2': {'before': Pt(10), 'after': Pt(8)},
    'Heading 3': {'before': Pt(10), 'after': Pt(8)},
    'Heading 4': {'before': Pt(10), 'after': Pt(8)},
}

TITLE_BODIES = {
    '设备用途', '性能特点', '机械结构', '自动控制系统', '系统操作及维护', '操作规程',
    '除砂机供风系统', '风阀', '气源处理系统', '电磁阀', '气缸',
    '除砂机机体', '结构与组成', '安装顺序和注意事项', '机体', '自动排料装置', '气动风阀系统',
    '开机及使用', '系统基本组成', '除砂机排料自动控制', '控制原理', '床层测量装置',
    '浮标密度的改变', '系统操作', '系统维护', '开车准备', '开车规定', '检查与调整',
    '停车规定', '除砂机维护保养',
    '工作原理', '使用与维护', '结构和组成', '可能出现的故障及排除方法',
}

STEP_STARTS = (
    '检查', '打开', '关闭', '开启', '停止', '启动', '接通', '调整', '确认', '观察',
    '首先', '然后', '其次', '最后', '待', '将', '按', '给料', '运转中', '运行中',
    '砂石床层形成后', '接到', '生产中发现', '定期', '清洗', '每两个月', '非电气',
    '正常情况下', '自动时', '合理地', '进入', '物料床层', '带料试验后',
)

SKIP_TEXTS = {
    'YTQS 系列除砂机使用说明书',
    '使用设备前，请详细阅读本说明书',
    '目 录',
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_run_font(run, cn_font, western_font, size_pt=None, bold=None, color=None):
    """Set Chinese and Western fonts for a single run."""
    run.font.name = western_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), cn_font)
    if size_pt:
        run.font.size = size_pt
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def apply_paragraph_font(p, cn_font, western_font, size_pt=None, bold=None, color=None):
    """Apply font settings to all runs in a paragraph."""
    for run in p.runs:
        set_run_font(run, cn_font, western_font, size_pt, bold, color)


def extract_prefix(text):
    """Extract leading numbering from a paragraph."""
    # Chinese numerals + 顿号
    m = re.match(r'^([一二三四五六七八九十百千万]+、)\s*', text)
    if m:
        return 'chinese', m.group(1), text[m.end():].strip()
    # Arabic decimal numbering, e.g. 1.1.1 / 1.1 / 1.
    m = re.match(r'^((?:\d+\.)+\d+|\d+\.)\s*', text)
    if m:
        return 'arabic_dot', m.group(1), text[m.end():].strip()
    # Single letter + space, e.g. a b c
    m = re.match(r'^([a-zA-Z])[\.、]\s*', text)
    if m:
        return 'alpha', m.group(1).lower(), text[m.end():].strip()
    return 'none', '', text


def heading_level_from_prefix(prefix):
    """Map an Arabic prefix like 1. / 1.1 / 1.1.1 to a Heading style."""
    if not prefix.endswith('.'):
        prefix += '.'
    dots = prefix.count('.') - 1
    if dots == 0:
        return 'Heading 2'
    elif dots == 1:
        return 'Heading 3'
    else:
        return 'Heading 4'


def is_title_text(text):
    """Conservative heuristic: does this numbered text look like a title?"""
    body = re.sub(r'^(\d+\.)+(\d+)?\s*', '', text).strip()
    body = re.sub(r'^[a-zA-Z][\.、]\s*', '', body)
    body = re.sub(r'^[一二三四五六七八九十]+、\s*', '', body)
    if not body:
        return True

    if body in TITLE_BODIES:
        return True
    if any(body.startswith(s) for s in STEP_STARTS):
        return False
    if '。' in body or '；' in body:
        return False
    if len(body) <= 12:
        return True
    if '，' not in body and '：' not in body and len(body) <= 18:
        return True
    return False


def style_heading(p, style_name):
    """Apply heading style (fonts, size, spacing, bold)."""
    p.style = style_name
    p.paragraph_format.space_before = HEADING_SPACING[style_name]['before']
    p.paragraph_format.space_after = HEADING_SPACING[style_name]['after']
    apply_paragraph_font(
        p,
        HEADING_CN_FONT,
        WESTERN_FONT,
        size_pt=HEADING_SIZES[style_name],
        bold=True,
        color=FONT_COLOR,
    )


def style_body(p, size_pt=Pt(12)):
    """Apply body text style."""
    p.style = 'Normal'
    apply_paragraph_font(p, BODY_CN_FONT, WESTERN_FONT, size_pt=size_pt)


def should_skip(text):
    """Skip headers/footers/cover text and very short artifacts."""
    if text in SKIP_TEXTS:
        return True
    if len(text) <= 2:
        return True
    if re.match(r'^\d+$', text):
        return True
    return False


# ---------------------------------------------------------------------------
# Main formatting logic
# ---------------------------------------------------------------------------

def format_document(doc):
    current_h2 = current_h3 = current_h4 = None
    step_counter = {}

    def reset_counters():
        step_counter.clear()

    def get_step_no(level_key):
        step_counter[level_key] = step_counter.get(level_key, 0) + 1
        return step_counter[level_key]

    def update_state(prefix):
        nonlocal current_h2, current_h3, current_h4
        parts = [p for p in prefix.rstrip('.').split('.') if p]
        current_h2 = parts[0] if len(parts) >= 1 else None
        current_h3 = f'{parts[0]}.{parts[1]}' if len(parts) >= 2 else None
        current_h4 = f'{parts[0]}.{parts[1]}.{parts[2]}' if len(parts) >= 3 else None
        reset_counters()

    # -----------------------------------------------------------------------
    # Pass 1: classify paragraphs and apply basic styles
    # -----------------------------------------------------------------------
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or should_skip(text):
            continue

        ptype, prefix, body = extract_prefix(text)

        if ptype == 'chinese':
            style_heading(p, 'Heading 1')
            current_h2 = current_h3 = current_h4 = None
            reset_counters()
            continue

        if ptype == 'arabic_dot':
            if is_title_text(text):
                style_name = heading_level_from_prefix(prefix)
                style_heading(p, style_name)
                update_state(prefix)
            else:
                level_key = current_h4 or current_h3 or current_h2 or 'root'
                no = get_step_no(level_key)
                p.text = f'{no}) {body}'
                style_body(p)
            continue

        if ptype == 'alpha':
            p.text = f'{prefix}) {body}'
            style_body(p)
            continue

        # Unnumbered paragraph — apply body font but leave style unchanged
        apply_paragraph_font(p, BODY_CN_FONT, WESTERN_FONT)

    # -----------------------------------------------------------------------
    # Pass 2: fixup known unnumbered headings
    # -----------------------------------------------------------------------
    current_h2 = current_h3 = current_h4 = None
    reset_counters()

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else 'Normal'

        if style.startswith('Heading'):
            ptype, prefix, body = extract_prefix(text)
            if ptype == 'chinese':
                current_h2 = current_h3 = current_h4 = None
                reset_counters()
            elif ptype == 'arabic_dot':
                update_state(prefix)
            continue

        # Add missing numbering to a few known headings
        if text == '结构和组成' and current_h3:
            p.text = f'{current_h3}.1 {text}'
            style_heading(p, 'Heading 4')
            update_state(f'{current_h3}.1')

    # -----------------------------------------------------------------------
    # Pass 3: format tables
    # -----------------------------------------------------------------------
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Horizontal alignment: left
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    apply_paragraph_font(paragraph, BODY_CN_FONT, WESTERN_FONT)

                # Vertical alignment: center
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                # Cell margins: top = 0.2 cm, left = 0.2 cm
                tcPr = cell._tc.get_or_add_tcPr()
                tcMar = tcPr.first_child_found_in('w:tcMar')
                if tcMar is None:
                    tcMar = OxmlElement('w:tcMar')
                    tcPr.append(tcMar)

                def set_margin(elem_name, value_cm):
                    elem = tcMar.find(qn(f'w:{elem_name}'))
                    if elem is None:
                        elem = OxmlElement(f'w:{elem_name}')
                        tcMar.append(elem)
                    elem.set(qn('w:w'), str(int(value_cm * 567)))
                    elem.set(qn('w:type'), 'dxa')

                set_margin('top', 0.2)
                set_margin('left', 0.2)

                # First row: bold + gray background
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'D9D9D9')
                    tcPr.append(shading_elm)

    return doc


def main():
    parser = argparse.ArgumentParser(description='Format a Chinese technical manual Word document.')
    parser.add_argument('input', help='Input .docx file path')
    parser.add_argument('output', help='Output .docx file path')
    args = parser.parse_args()

    doc = Document(args.input)
    format_document(doc)
    doc.save(args.output)
    print(f'Formatted document saved to: {args.output}')


if __name__ == '__main__':
    main()
